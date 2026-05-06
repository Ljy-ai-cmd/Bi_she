#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('示例文档', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('这是一个使用 python-docx 库生成的 Word 文档。')

doc.add_heading('功能特点', level=2)

doc.add_paragraph('• 支持多种文本格式', style='List Bullet')
doc.add_paragraph('• 支持标题和段落', style='List Bullet')
doc.add_paragraph('• 支持列表', style='List Bullet')
doc.add_paragraph('• 支持表格', style='List Bullet')

doc.add_heading('示例表格', level=2)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '序号'
hdr_cells[1].text = '名称'
hdr_cells[2].text = '描述'

row1_cells = table.rows[1].cells
row1_cells[0].text = '1'
row1_cells[1].text = '项目A'
row1_cells[2].text = '这是第一个项目'

row2_cells = table.rows[2].cells
row2_cells[0].text = '2'
row2_cells[1].text = '项目B'
row2_cells[2].text = '这是第二个项目'

doc.add_paragraph()
doc.add_paragraph('文档创建时间: 2026-05-05')
doc.add_paragraph('—— 结束 ——').alignment = WD_ALIGN_PARAGRAPH.CENTER

output_path = 'e:/biyesheji/SUTrack-main11/示例文档.docx'
doc.save(output_path)
print(f'Word 文档已创建: {output_path}')
